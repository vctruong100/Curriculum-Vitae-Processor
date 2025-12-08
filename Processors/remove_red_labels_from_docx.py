#!/usr/bin/env python3
import argparse
import csv
import os
import re
import string
from typing import Dict, List, Tuple, Optional

YEAR_RE = re.compile(r'^\s*(\d{4})\b')
PHASE_RE = re.compile(r'^\s*phase(\s|$)', re.IGNORECASE)

DEFAULT_SECTION_START = "Research Experience"
DEFAULT_SECTION_END   = "By signing this form, I confirm that the information provided is accurate and reflects my current qualifications."

try:
    from docx import Document
except Exception:
    Document = None

# ---------- Normalization & similarity ----------

def _norm_ws(s: str) -> str:
    return ' '.join((s or '').replace('\t',' ').split())

def _after_year(text: str) -> str:
    m = YEAR_RE.match(text or '')
    if not m:
        return _norm_ws(text or '')
    return _norm_ws((text or '')[m.end(0):])

def _normalize_for_match(s: str) -> str:
    # Lower, remove punctuation (but keep spaces), collapse spaces
    s = (s or '').lower()
    s = s.translate(str.maketrans({c:' ' for c in string.punctuation}))
    return _norm_ws(s)

def _similarity(a: str, b: str) -> float:
    # SequenceMatcher on normalized strings
    try:
        import difflib
        return difflib.SequenceMatcher(None, _normalize_for_match(a), _normalize_for_match(b)).ratio()
    except Exception:
        # Fallback: token overlap
        A = set(_normalize_for_match(a).split())
        B = set(_normalize_for_match(b).split())
        if not A and not B: return 1.0
        if not A or not B:  return 0.0
        return len(A & B) / len(A | B)

# ---------- Bounds for CV ----------

def _find_bounds_paras(doc, start_text: str, end_text: str):
    def norm(s): return ' '.join((s or '').strip().lower().split())
    paras = list(doc.paragraphs)
    s_norm = norm(start_text); e_norm = norm(end_text)
    s_idx = None
    for i, p in enumerate(paras):
        if s_norm in norm(p.text):
            s_idx = i; break
    if s_idx is None:
        raise RuntimeError(f'Could not find section start containing: "{start_text}".')
    e_idx = None
    for j in range(s_idx + 1, len(paras)):
        if e_norm in norm(paras[j].text):
            e_idx = j; break
    if e_idx is None: e_idx = len(paras)
    return paras, s_idx, e_idx

def _is_year_paragraph_text(text: str) -> bool:
    return YEAR_RE.match(text or "") is not None

def _clear_runs(p) -> None:
    for r in list(p.runs):
        r.clear()
        p._element.remove(r._element)

def _write_study_paragraph(p, year: str, after_text: str) -> None:
    """Write YEAR + <tab> + after_text; bold until first colon (if present)."""
    _clear_runs(p)
    p.add_run(year).bold = False
    p.add_run('\t')  # real tab
    colon_pos = after_text.find(':')
    if colon_pos == -1:
        p.add_run(after_text)
        return
    before = after_text[:colon_pos]
    after  = after_text[colon_pos:]  # includes ':'
    r1 = p.add_run(before); r1.bold = True
    p.add_run(after)

# ---------- Category pruning (no-studies) ----------

def _is_category_paragraph(p) -> bool:
    txt = (p.text or '').strip()
    if not txt or _is_year_paragraph_text(txt):
        return False
    try:
        if p.style and 'heading' in p.style.name.lower():
            return True
    except Exception:
        pass
    # Heuristics: no digits, reasonably short, colon/uppercase/bold hint
    if any(ch.isdigit() for ch in txt):
        return False
    if len(txt) <= 80:
        ends_colon = txt.endswith(':')
        mostly_upper = sum(c.isupper() for c in txt if c.isalpha()) >= 0.6 * sum(1 for c in txt if c.isalpha())
        all_bold = (len(p.runs) > 0) and all((r.font.bold is True) for r in p.runs if r.text.strip())
        if ends_colon or mostly_upper or all_bold:
            return True
    return False

def _delete_paragraph(p):
    p._element.getparent().remove(p._element)

def _prune_empty_categories(doc, section_start: str, section_end: str):
    """Remove category paragraphs inside Research Experience that have no studies until next category/end.
       NEVER remove Phase headers (e.g., 'Phase I', 'Phase II-IV').
    """
    paras, s_idx, e_idx = _find_bounds_paras(doc, section_start, section_end)
    start = s_idx + 1
    if start >= e_idx:
        return
    block = paras[start:e_idx]

    # Collect indices of candidate categories
    cat_global_idxs = []
    for off, p in enumerate(block):
        if _is_category_paragraph(p):
            cat_global_idxs.append(start + off)

    to_delete = []
    for k, gi in enumerate(cat_global_idxs):
        p_cat = paras[gi]
        txt = (p_cat.text or '').strip()
        if PHASE_RE.match(txt):  # keep Phase headers
            continue
        scan_from = gi + 1
        scan_to   = cat_global_idxs[k+1] if k+1 < len(cat_global_idxs) else e_idx
        has_study = False
        for j in range(scan_from, scan_to):
            if _is_year_paragraph_text(paras[j].text or ''):
                has_study = True
                break
        if not has_study:
            to_delete.append(gi)

    for di in sorted(to_delete, reverse=True):
        _delete_paragraph(paras[di])

# ---------- CSV mapping ----------

def load_mapping_by_year(path: str) -> Dict[str, List[Tuple[str, str]]]:
    """
    Returns dict year -> list of (red_after, nonred_after).
    CSV columns: year, red_study_text, nonred_study_text.
    Study text may include or omit the year; we strip it for matching.
    Robust to encodings: tries utf-8-sig first, then latin-1.
    """
    by_year: Dict[str, List[Tuple[str, str]]] = {}
    encs = ['utf-8-sig', 'latin-1']
    last_err = None
    for enc in encs:
        try:
            with open(path, 'r', encoding=enc, newline='') as f:
                reader = csv.reader(f)
                for row in reader:
                    if not row or len(row) < 3:
                        continue
                    year = str(row[0]).strip()
                    if not (year.isdigit() and len(year) == 4):
                        # header or invalid
                        continue
                    red_after    = _after_year(str(row[1]))
                    nonred_after = _after_year(str(row[2]))
                    by_year.setdefault(year, []).append((red_after, nonred_after))
            return by_year
        except UnicodeDecodeError as e:
            last_err = e
            continue
    raise RuntimeError(f"Could not read mapping CSV with utf-8 or latin-1: {last_err}")

# ---------- Main process ----------

def process_fuzzy_csv(original_cv: str, mapping_csv: str, out_path: Optional[str],
                      section_start: str, section_end: str, threshold: float) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required.")
    if not os.path.isfile(original_cv):
        raise FileNotFoundError(f"Original CV not found: {original_cv}")
    if not os.path.isfile(mapping_csv):
        raise FileNotFoundError(f"Mapping CSV not found: {mapping_csv}")

    by_year = load_mapping_by_year(mapping_csv)

    from shutil import copyfile
    final_out = out_path or original_cv
    try:
        if os.path.abspath(final_out) != os.path.abspath(original_cv):
            copyfile(original_cv, final_out)
    except Exception:
        raise PermissionError("Permission denied while writing the output. Please close ALL CV documents.")

    doc = Document(final_out)

    # CV bounds; if not found, scan whole doc
    try:
        paras, s_idx, e_idx = _find_bounds_paras(doc, section_start, section_end)
        scan_range = range(s_idx + 1, e_idx)
    except RuntimeError:
        paras = list(doc.paragraphs)
        scan_range = range(0, len(paras))

    changed = 0
    attempted = 0

    for i in scan_range:
        p = paras[i]
        txt = p.text or ''
        if not _is_year_paragraph_text(txt):
            continue

        m = YEAR_RE.match(txt)
        if not m:
            continue
        year = m.group(1)
        after_cv = _after_year(txt)

        candidates = by_year.get(year, [])
        if not candidates:
            continue

        # Find best fuzzy match among CSV col2 (red_after)
        best_idx = -1
        best_score = -1.0
        for idx, (red_after, nonred_after) in enumerate(candidates):
            s = _similarity(after_cv, red_after)
            if s > best_score:
                best_score = s
                best_idx = idx

        if best_idx == -1 or best_score < threshold:
            continue

        attempted += 1
        _, replacement_after = candidates[best_idx]
        _write_study_paragraph(p, year, replacement_after)
        changed += 1

    # Prune empty categories after replacements
    try:
        _prune_empty_categories(doc, section_start, section_end)
    except Exception:
        pass

    doc.save(final_out)
    return f"Attempted: {attempted}, Replaced: {changed}, Output: {final_out}"

def main():
    ap = argparse.ArgumentParser(description="Remove red labels using a CSV mapping with fuzzy matching; prune empty categories; no resorting.")
    ap.add_argument('--original-cv', required=True, help='Path to the FINAL CV (.docx)')
    ap.add_argument('--mapping-csv', required=True, help='Path to CSV: year, red_study_text, nonred_study_text')
    ap.add_argument('--out', default=None, help='Path to output .docx (defaults to overwrite original)')
    ap.add_argument('--section-start', default=DEFAULT_SECTION_START, help='CV section start marker (if present)')
    ap.add_argument('--section-end',   default=DEFAULT_SECTION_END,   help='CV section end marker (if present)')
    ap.add_argument('--threshold', type=float, default=0.90, help='Fuzzy match threshold (0–1). Default 0.90')
    args = ap.parse_args()

    msg = process_fuzzy_csv(args.original_cv, args.mapping_csv, args.out, args.section_start, args.section_end, args.threshold)
    print(msg)

if __name__ == "__main__":
    main()
