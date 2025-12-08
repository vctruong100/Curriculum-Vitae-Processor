#!/usr/bin/env python3
"""
Phase-aware sorter that categorizes an unsorted study list using a MASTER text list,
**but outputs the MASTER-formatted study text** placed under the correct Phase and Category.

Fixes covered:
  - Categories with the same name can appear under multiple phases (Phase I and Phase II-IV).
    Studies are placed in the specific Phase->Category from the MASTER entry that was matched.
  - Text output does NOT carry over '**' markers by default (toggle with --text-bold-markers if desired).
  - Per-category ordering is **descending by year** (e.g., 2024, 2023, 2022...).
  - DOCX output bolds only the protocol name (text before the first ':') and uses a
    **hanging indent + tab stop** so natural wraps align under the title.

Usage example:
  python sorterv3_masterfmt_phase.py \
    --master master_study_list.txt \
    --unsorted unsorted_study_list.txt \
    --out sorted_unsorted_study_list.txt \
    --audit match_audit_report.txt \
    --threshold 0.80 \
    --docx-out sorted_unsorted_study_list.docx \
    --indent-type spaces --indent-size 1 \
    --docx-indent 0.5 \
    --text-bold-markers false --bold true
"""
import argparse
import difflib
import os
import re
import string
from collections import OrderedDict
from typing import List, Tuple, Optional, Dict

try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception:
    Document = None
    Pt = None
    Inches = None

# -----------------
# Regexes & constants
# -----------------
YEAR_LINE = re.compile(r'^\s*(\d{4})\b')
MULTI_X = re.compile(r'x{3,}', flags=re.IGNORECASE)

# Phase headers: accept variants and punctuation
PHASE_I_PATTERNS = [
    r'^\s*phase\s*i\s*[:.\-–—]?\s*$',
    r'^\s*phase\s*1\s*[:.\-–—]?\s*$'
]
PHASE_II_IV_PATTERNS = [
    r'^\s*phase\s*ii\s*[-–—/ ]*iv\s*[:.\-–—]?\s*$',
    r'^\s*phase\s*2\s*[-–—/ ]*4\s*[:.\-–—]?\s*$'
]
PHASE_I_LABEL = 'Phase I'
PHASE_II_IV_LABEL = 'Phase II-IV'
PHASE_I_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_I_PATTERNS]
PHASE_II_IV_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_II_IV_PATTERNS]

# Category bullet acceptance
CATEGORY_BULLET_RE = re.compile(r'^\s*[-–—•●]\s*(.+?)\s*:?\s*$')

# -----------------
# Helpers
# -----------------

def clean_space_tabs(text: str) -> str:
    text = text.replace('\t', ' ')
    text = ' '.join(text.split())
    return text


def extract_year(text: str) -> Optional[str]:
    m = YEAR_LINE.match(text)
    if m is not None:
        return m.group(1)
    return None


def strip_leading_year(text: str) -> str:
    if YEAR_LINE.match(text):
        return re.sub(r'^\s*\d{4}\s*', '', text)
    return text


def is_phase_header(line: str) -> Optional[str]:
    stripped = line.strip()
    for r in PHASE_I_RE:
        if r.match(stripped):
            return PHASE_I_LABEL
    for r in PHASE_II_IV_RE:
        if r.match(stripped):
            return PHASE_II_IV_LABEL
    return None


def normalize_text(text: str) -> str:
    # Lower, remove leading year, remove long x placeholders, strip punctuation including '*'
    t = text.lower()
    t = re.sub(r'^\s*\d{4}\s*', '', t)
    t = MULTI_X.sub('', t)
    # Remove punctuation including asterisks used for markdown bold markers
    t = t.translate(str.maketrans('', '', string.punctuation))
    t = ' '.join(t.split())
    return t


def token_set(text: str) -> set:
    stop = {
        'a','an','the','to','of','and','in','for','with','on','by','from',
        'study','phase','randomized','openlabel','double','single','dose',
        'multiple','ascending','participants','healthy','placebo','controlled',
        'evaluate','assess','safety','tolerability'
    }
    tokens = normalize_text(text).split()
    filtered: List[str] = []
    for t in tokens:
        if t in stop:
            continue
        filtered.append(t)
    return set(filtered)


def jaccard_similarity(a: str, b: str) -> float:
    sa = token_set(a)
    sb = token_set(b)
    if len(sa) == 0 and len(sb) == 0:
        return 1.0
    if len(sa) == 0 or len(sb) == 0:
        return 0.0
    inter = len(sa & sb)
    union = len(sa | sb)
    if union == 0:
        return 0.0
    return inter / union


def combined_similarity(a: str, b: str, w_seq: float = 0.6, w_jac: float = 0.4) -> float:
    seq = difflib.SequenceMatcher(None, normalize_text(a), normalize_text(b)).ratio()
    jac = jaccard_similarity(a, b)
    score = (w_seq * seq) + (w_jac * jac)
    return score

# -----------------
# Parsing MASTER (Phase -> Category -> Studies)
# -----------------

def parse_master_hierarchy(master_path: str) -> "OrderedDict[str, OrderedDict[str, List[str]]]":
    phases: "OrderedDict[str, OrderedDict[str, List[str]]]" = OrderedDict()
    current_phase: Optional[str] = None
    current_category: Optional[str] = None
    current_study_buffer: Optional[str] = None

    def commit_study() -> None:
        nonlocal current_study_buffer
        if current_study_buffer is not None and current_phase is not None and current_category is not None:
            study_text = clean_space_tabs(current_study_buffer)
            phases[current_phase][current_category].append(study_text)
            current_study_buffer = None

    with open(master_path, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')

            if len(line.strip()) == 0:
                commit_study()
                continue

            # Phase header takes precedence
            ph = is_phase_header(line)
            if ph is not None:
                commit_study()
                current_phase = ph
                if current_phase not in phases:
                    phases[current_phase] = OrderedDict()
                current_category = None
                continue

            # New study begins?
            if YEAR_LINE.match(line):
                commit_study()
                if current_phase is None:
                    current_phase = PHASE_I_LABEL
                    if current_phase not in phases:
                        phases[current_phase] = OrderedDict()
                if current_category is None:
                    current_category = 'Uncategorized (from master file)'
                    if current_category not in phases[current_phase]:
                        phases[current_phase][current_category] = []
                current_study_buffer = line.strip()
                continue

            # If we are inside a study, treat subsequent lines as continuation
            if current_study_buffer is not None:
                current_study_buffer = f"{current_study_buffer} {line.strip()}"
                continue

            # Category header
            m = CATEGORY_BULLET_RE.match(line)
            if m is not None:
                cat = clean_space_tabs(m.group(1))
            else:
                stripped = line.strip()
                if stripped.endswith(':'):
                    stripped = stripped[:-1].strip()
                cat = clean_space_tabs(stripped)

            if current_phase is None:
                current_phase = PHASE_I_LABEL
                if current_phase not in phases:
                    phases[current_phase] = OrderedDict()
            current_category = cat
            if current_category not in phases[current_phase]:
                phases[current_phase][current_category] = []

    commit_study()
    return phases

# -----------------
# Parsing UNSORTED studies list (flat)
# -----------------

def parse_unsorted_studies(unsorted_path: str) -> List[str]:
    studies: List[str] = []
    current: Optional[str] = None

    with open(unsorted_path, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')

            if len(line.strip()) == 0:
                if current is not None:
                    studies.append(clean_space_tabs(current))
                    current = None
                continue

            if YEAR_LINE.match(line):
                if current is not None:
                    studies.append(clean_space_tabs(current))
                current = line.strip()
            else:
                if current is None:
                    continue
                current = f"{current} {line.strip()}"

    if current is not None:
        studies.append(clean_space_tabs(current))

    return studies

# -----------------
# Build index over MASTER studies
# -----------------

def build_master_index(phases: "OrderedDict[str, OrderedDict[str, List[str]]]") -> Tuple[List[str], List[str], List[str], List[str], List[Optional[str]]]:
    m_orig: List[str] = []
    m_norm: List[str] = []
    m_phase: List[str] = []
    m_cat: List[str] = []
    m_year: List[Optional[str]] = []

    for phase, cats in phases.items():
        for cat, studies in cats.items():
            for s in studies:
                m_orig.append(s)
                m_norm.append(normalize_text(s))
                m_phase.append(phase)
                m_cat.append(cat)
                m_year.append(extract_year(s))

    return m_orig, m_norm, m_phase, m_cat, m_year

# -----------------
# Protocol name split (from MASTER text after year)
# -----------------

def split_name_protocol(line_after_year: str) -> Tuple[str, bool, str]:
    idx = line_after_year.find(':')
    if idx == -1:
        return line_after_year.strip(), False, ''
    name = line_after_year[:idx].strip()
    remainder = line_after_year[idx+1:].lstrip()
    return name, True, remainder

# -----------------
# Categorization (match unsorted to MASTER, then place under MASTER phase+category)
# -----------------

def categorize_with_master(phases: "OrderedDict[str, OrderedDict[str, List[str]]]",
                           unsorted_studies: List[str],
                           threshold: float = 0.80
                           ) -> Tuple["OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]", List[Tuple[str, Optional[str], Optional[str], Optional[str], float]]]:
    """Return:
      - categorized: OrderedDict[phase][category] -> list of (year, master_after_year)
      - audit_rows: (unsorted_text, matched_master_or_None, phase_or_None, category_or_None, score)
    We always output the MASTER text (after year), not the unsorted string.
    """
    categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]" = OrderedDict()
    for ph, cats in phases.items():
        categorized[ph] = OrderedDict()
        for cat in cats.keys():
            categorized[ph][cat] = []
    if 'Uncategorized' not in categorized:
        categorized['Uncategorized'] = OrderedDict()
    if 'Uncategorized' not in categorized['Uncategorized']:
        categorized['Uncategorized']['Uncategorized'] = []

    m_orig, m_norm, m_phase, m_cat, m_year = build_master_index(phases)

    # Exact normalized map across all phases for the same year
    exact_map: Dict[Tuple[Optional[str], str], List[int]] = {}
    # Year filter to candidate indices
    year_to_indices: Dict[Optional[str], List[int]] = {}

    for i, nn in enumerate(m_norm):
        key = (m_year[i], nn)
        if key not in exact_map:
            exact_map[key] = [i]
        else:
            exact_map[key].append(i)

        y = m_year[i]
        if y not in year_to_indices:
            year_to_indices[y] = [i]
        else:
            year_to_indices[y].append(i)

    audit_rows: List[Tuple[str, Optional[str], Optional[str], Optional[str], float]] = []

    for u in unsorted_studies:
        u_year = extract_year(u)
        u_norm = normalize_text(u)

        # 1) Try exact across phases
        key = (u_year, u_norm)
        if key in exact_map:
            idx = exact_map[key][0]
            ph = m_phase[idx]
            cat = m_cat[idx]
            master_line = m_orig[idx]
            categorized[ph][cat].append((u_year or extract_year(master_line) or '', strip_leading_year(master_line).lstrip()))
            audit_rows.append((u, master_line, ph, cat, 1.0))
            continue

        # 2) Fuzzy within same year across phases
        best_idx: Optional[int] = None
        best_score: float = -1.0
        candidates = year_to_indices.get(u_year, [])

        if len(candidates) == 0:
            categorized['Uncategorized']['Uncategorized'].append((u_year or '', strip_leading_year(u)))
            audit_rows.append((u, None, None, None, 0.0))
            continue

        for i in candidates:
            sc = combined_similarity(u, m_orig[i])
            if sc > best_score:
                best_score = sc
                best_idx = i

        if best_idx is not None and best_score >= threshold:
            ph = m_phase[best_idx]
            cat = m_cat[best_idx]
            master_line = m_orig[best_idx]
            categorized[ph][cat].append((u_year or extract_year(master_line) or '', strip_leading_year(master_line).lstrip()))
            audit_rows.append((u, master_line, ph, cat, best_score))
        else:
            categorized['Uncategorized']['Uncategorized'].append((u_year or '', strip_leading_year(u)))
            safe = best_score if best_idx is not None else 0.0
            audit_rows.append((u, None, None, None, safe))

    # 3) Sort each Phase->Category by year desc
    def year_as_int(yy: str) -> int:
        try:
            return int(yy)
        except Exception:
            return -1

    for ph, cats in categorized.items():
        for cat, items in cats.items():
            items_sorted = sorted(items, key=lambda t: year_as_int(t[0]), reverse=True)
            categorized[ph][cat] = items_sorted

    return categorized, audit_rows

# -----------------
# Writers
# -----------------

def write_text(out_path: str,
               categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]",
               indent_sep: str = ' ',
               bold_markers: bool = False,
               bold: bool = True) -> None:
    with open(out_path, 'w', encoding='utf-8') as f:
        for ph, cats in categorized.items():
            f.write(f"{ph}\n")
            for cat, items in cats.items():
                f.write(f"{cat}\n")
                for (year, master_after_year) in items:
                    name, has_colon, remainder = split_name_protocol(master_after_year)
                    if bold and bold_markers:
                        first = f"{year}{indent_sep}**{name}**"
                    else:
                        first = f"{year}{indent_sep}{name}"
                    if has_colon:
                        if len(remainder) > 0:
                            first = f"{first}: {remainder}"
                        else:
                            first = f"{first}:"
                    f.write(first + "\n")
                f.write("\n")
            f.write("\n")


def set_hanging_indent_with_tab(p, indent_inch: float = 0.5) -> None:
    if Inches is None:
        return
    pf = p.paragraph_format
    indent = Inches(indent_inch)
    pf.left_indent = indent
    pf.first_line_indent = -indent
    try:
        pf.tab_stops.add_tab_stop(indent)
    except Exception:
        pass


def write_docx(docx_path: str,
               categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]",
               docx_indent_inch: float = 0.5,
               bold: bool = True) -> None:
    if Document is None:
        raise RuntimeError('python-docx is not available.')

    doc = Document()

    for ph, cats in categorized.items():
        p_phase = doc.add_paragraph(ph)
        for run in p_phase.runs:
            run.bold = True

        for cat, items in cats.items():
            p_cat = doc.add_paragraph(cat)
            for run in p_cat.runs:
                run.bold = True

            for (year, master_after_year) in items:
                p = doc.add_paragraph()
                set_hanging_indent_with_tab(p, indent_inch=docx_indent_inch)

                # YEAR + TAB
                r_year = p.add_run(year)
                r_year.bold = False
                p.add_run('\t')

                name, has_colon, remainder = split_name_protocol(master_after_year)

                # Bold protocol name only
                r_name = p.add_run(name)
                r_name.bold = bool(bold)

                if has_colon:
                    p.add_run(':')
                    if len(remainder) > 0:
                        p.add_run(' ' + remainder)

    doc.save(docx_path)

# -----------------
# Entry point
# -----------------
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Phase-aware sorter that outputs MASTER-formatted text by phase/category and year desc.')
    parser.add_argument('--master', required=True, help='MASTER text file (with Phase headers, categories, and year-starting studies)')
    parser.add_argument('--unsorted', required=True, help='Unsorted studies text file (year-starting lines + continuations)')
    parser.add_argument('--out', required=True, help='Text output path')
    parser.add_argument('--audit', required=True, help='Audit TSV path')
    parser.add_argument('--threshold', type=float, default=0.80)

    # Text formatting
    parser.add_argument('--indent-type', choices=['tab', 'spaces'], default='spaces')
    parser.add_argument('--indent-size', type=int, default=1)
    parser.add_argument('--bold', type=str, default='true')
    parser.add_argument('--text-bold-markers', type=str, default='false', help='If true, wrap protocol name with **..** in TEXT output')

    # DOCX output
    parser.add_argument('--docx-out', default=None)
    parser.add_argument('--docx-indent', type=float, default=0.5)

    args = parser.parse_args()

    # Resolve indent separator for TEXT output (DOCX uses tab + hanging indent)
    if args.indent_type == 'tab':
        indent_sep = '\t'
    else:
        if args.indent_size < 1:
            indent_sep = ' '
        else:
            indent_sep = ' ' * args.indent_size

    def to_bool(val: str) -> bool:
        v = str(val).strip().lower()
        if v in ('1','true','yes','y'):
            return True
        if v in ('0','false','no','n'):
            return False
        return False

    bold_flag = to_bool(args.bold)
    bold_markers_flag = to_bool(args.text_bold_markers)

    if not os.path.isfile(args.master):
        print(f'ERROR: Master file not found: {args.master}')
        raise SystemExit(2)
    if not os.path.isfile(args.unsorted):
        print(f'ERROR: Unsorted file not found: {args.unsorted}')
        raise SystemExit(2)

    try:
        master_phases = parse_master_hierarchy(args.master)
        unsorted_list = parse_unsorted_studies(args.unsorted)
        categorized, audit = categorize_with_master(master_phases, unsorted_list, args.threshold)

        write_text(args.out, categorized, indent_sep=indent_sep, bold_markers=bold_markers_flag, bold=bold_flag)

        if args.docx_out is not None:
            write_docx(args.docx_out, categorized, docx_indent_inch=args.docx_indent, bold=bold_flag)

        with open(args.audit, 'w', encoding='utf-8') as f:
            f.write('Unsorted\tMatched_Master\tPhase\tCategory\tScore\n')
            for u, m, ph, cat, sc in audit:
                f.write(f"{u}\t{m or ''}\t{ph or 'Uncategorized'}\t{cat or 'Uncategorized'}\t{sc:.3f}\n")

        print('Done.')
    except Exception as e:
        print(f'ERROR: {e}')
        raise SystemExit(1)
