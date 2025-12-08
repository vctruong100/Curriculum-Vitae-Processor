#!/usr/bin/env python3
"""
Merge existing studies from a **sorted .docx** (not .txt) with NEW studies from a master .docx,
then output a full .docx containing ALL previous studies + the newly inserted ones.

- Reads the existing sorted .docx (the one you generated previously) to get:
  - Phase -> Category -> Studies (as Word runs)
  - The latest year across all categories
- Reads the master .docx to parse Phase -> Category -> Studies (as runs) that preserve RED labeling
- Adds only those master studies whose year is strictly GREATER than the latest year found in the sorted .docx
- Writes a single merged .docx with both previous + new studies
- Keeps the run formatting of previous docx as-is; for new ones, preserves red labels from master
- Uses hanging indent + tab so natural wraps align under the title

Usage:
  python merge_insert_red_docx_v2.py \
    --existing-docx sorted_unsorted_study_list.docx \
    --master-docx master_study_list.docx \
    --out-docx merged_with_new_studies.docx \
    --indent 0.5
"""
import argparse
import os
import re
from collections import OrderedDict
from typing import List, Optional, Tuple, Any

try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_BREAK
    from docx.enum.dml import MSO_THEME_COLOR
except Exception:
    Document = None
    Inches = None
    RGBColor = None
    WD_BREAK = None
    MSO_THEME_COLOR = None

YEAR_LINE = re.compile(r'^\s*(\d{4})\b')
PHASE_I_PATTERNS = [r'^\s*phase\s*i\s*[:.\-–—]?\s*$', r'^\s*phase\s*1\s*[:.\-–—]?\s*$']
PHASE_II_IV_PATTERNS = [r'^\s*phase\s*ii\s*[-–—/ ]*iv\s*[:.\-–—]?\s*$', r'^\s*phase\s*2\s*[-–—/ ]*4\s*[:.\-–—]?\s*$']
PHASE_I_LABEL = 'Phase I'
PHASE_II_IV_LABEL = 'Phase II-IV'
PHASE_I_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_I_PATTERNS]
PHASE_II_IV_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_II_IV_PATTERNS]


def is_phase_header(text: str) -> Optional[str]:
    stripped = text.strip()
    for r in PHASE_I_RE:
        if r.match(stripped):
            return PHASE_I_LABEL
    for r in PHASE_II_IV_RE:
        if r.match(stripped):
            return PHASE_II_IV_LABEL
    return None


class RunSeg:
    def __init__(self, text: str, bold: bool, rgb: Optional[Tuple[int,int,int]], theme):
        self.text = text
        self.bold = bold
        self.rgb = rgb
        self.theme = theme


class StudyRuns:
    def __init__(self, year: str, runs: List[RunSeg]):
        self.year = year
        self.runs = runs  # runs AFTER the year (paragraph starts with year)


def get_runs(p) -> List[RunSeg]:
    segs: List[RunSeg] = []
    for r in p.runs:
        t = r.text
        if not t:
            continue
        rgb = None
        theme = None
        try:
            if r.font is not None and r.font.color is not None:
                if r.font.color.rgb is not None:
                    val = r.font.color.rgb
                    rgb = (val[0], val[1], val[2])
                elif r.font.color.theme_color is not None:
                    theme = r.font.color.theme_color
        except Exception:
            pass
        segs.append(RunSeg(t, bool(r.bold), rgb, theme))
    return segs


def split_off_year(runs: List[RunSeg]) -> Tuple[str, List[RunSeg]]:
    full = ''.join(s.text for s in runs)
    m = YEAR_LINE.match(full)
    if not m:
        return '', runs
    year = m.group(1)
    cut = m.end(1)
    while cut < len(full) and full[cut] in (' ', '\t'):
        cut += 1

    # slice runs by char count
    rem: List[RunSeg] = []
    consumed = 0
    for s in runs:
        L = len(s.text)
        if consumed + L <= cut:
            consumed += L
            continue
        if consumed < cut < consumed + L:
            offset = cut - consumed
            rem.append(RunSeg(s.text[offset:], s.bold, s.rgb, s.theme))
            consumed += L
            continue
        rem.append(s)
        consumed += L
    return year, rem


def parse_studies_from_docx(docx_path: str) -> "OrderedDict[str, OrderedDict[str, List[StudyRuns]]]":
    if Document is None:
        raise RuntimeError('python-docx is required')
    doc = Document(docx_path)
    phases: "OrderedDict[str, OrderedDict[str, List[StudyRuns]]]" = OrderedDict()
    cur_phase = None
    cur_cat = None

    for p in doc.paragraphs:
        txt = p.text or ''
        if len(txt.strip()) == 0:
            continue
        ph = is_phase_header(txt)
        if ph is not None:
            cur_phase = ph
            if cur_phase not in phases:
                phases[cur_phase] = OrderedDict()
            cur_cat = None
            continue
        if YEAR_LINE.match(txt):
            if cur_phase is None:
                cur_phase = PHASE_I_LABEL
                if cur_phase not in phases:
                    phases[cur_phase] = OrderedDict()
            if cur_cat is None:
                cur_cat = 'Uncategorized'
                if cur_cat not in phases[cur_phase]:
                    phases[cur_phase][cur_cat] = []
            runs = get_runs(p)
            y, rem = split_off_year(runs)
            if not y:
                m = YEAR_LINE.match(txt)
                y = m.group(1) if m else '0000'
            phases[cur_phase][cur_cat].append(StudyRuns(y, rem))
            continue
        # else category
        cur_cat = txt.strip()
        if cur_phase is None:
            cur_phase = PHASE_I_LABEL
            if cur_phase not in phases:
                phases[cur_phase] = OrderedDict()
        if cur_cat not in phases[cur_phase]:
            phases[cur_phase][cur_cat] = []

    return phases


def latest_year_from_phases(phases) -> Optional[int]:
    latest = None
    for ph, cats in phases.items():
        for cat, lst in cats.items():
            for st in lst:
                try:
                    y = int(st.year)
                except Exception:
                    y = 0
                if latest is None or y > latest:
                    latest = y
    return latest


def set_hanging_indent_with_tab(p, indent_inch: float = 0.5):
    if Inches is None:
        return
    pf = p.paragraph_format
    ind = Inches(indent_inch)
    pf.left_indent = ind
    pf.first_line_indent = -ind
    try:
        pf.tab_stops.add_tab_stop(ind)
    except Exception:
        pass


def write_runs(p, runs: List[RunSeg], bold_until_colon: bool = True):
    before_colon = True
    for s in runs:
        text = s.text
        i = 0
        while i < len(text):
            if before_colon:
                pos = text.find(':', i)
                if pos == -1:
                    chunk = text[i:]
                    r = p.add_run(chunk)
                    r.bold = bool(s.bold) or (bold_until_colon and len(chunk) > 0)
                    try:
                        if s.rgb is not None and RGBColor is not None:
                            r.font.color.rgb = RGBColor(s.rgb[0], s.rgb[1], s.rgb[2])
                        elif s.theme is not None and MSO_THEME_COLOR is not None:
                            r.font.color.theme_color = s.theme
                    except Exception:
                        pass
                    break
                else:
                    chunk = text[i:pos]
                    if chunk:
                        r = p.add_run(chunk)
                        r.bold = bool(s.bold) or True  # ensure bold before colon
                        try:
                            if s.rgb is not None and RGBColor is not None:
                                r.font.color.rgb = RGBColor(s.rgb[0], s.rgb[1], s.rgb[2])
                            elif s.theme is not None and MSO_THEME_COLOR is not None:
                                r.font.color.theme_color = s.theme
                        except Exception:
                            pass
                    r2 = p.add_run(':')
                    try:
                        if s.rgb is not None and RGBColor is not None:
                            r2.font.color.rgb = RGBColor(s.rgb[0], s.rgb[1], s.rgb[2])
                        elif s.theme is not None and MSO_THEME_COLOR is not None:
                            r2.font.color.theme_color = s.theme
                    except Exception:
                        pass
                    before_colon = False
                    i = pos + 1
            else:
                chunk = text[i:]
                r = p.add_run(chunk)
                r.bold = bool(s.bold) and not bold_until_colon
                try:
                    if s.rgb is not None and RGBColor is not None:
                        r.font.color.rgb = RGBColor(s.rgb[0], s.rgb[1], s.rgb[2])
                    elif s.theme is not None and MSO_THEME_COLOR is not None:
                        r.font.color.theme_color = s.theme
                except Exception:
                    pass
                break


def merge_write(existing_docx: str, master_docx: str, out_docx: str, indent_inch: float = 0.5):
    if Document is None:
        raise RuntimeError('python-docx is required')

    exist_phases = parse_studies_from_docx(existing_docx)
    latest = latest_year_from_phases(exist_phases)

    master_phases = parse_studies_from_docx(master_docx)

    # Build combined: start with all existing
    combined: "OrderedDict[str, OrderedDict[str, List[Tuple[str, StudyRuns]]]]" = OrderedDict()
    for ph, cats in exist_phases.items():
        combined[ph] = OrderedDict()
        for cat, lst in cats.items():
            combined[ph][cat] = [('old', st) for st in lst]

    # Append new ones from master with year > latest
    for ph, cats in master_phases.items():
        if ph not in combined:
            combined[ph] = OrderedDict()
        for cat, lst in cats.items():
            for st in lst:
                try:
                    y = int(st.year)
                except Exception:
                    y = 0
                if latest is not None and y <= latest:
                    continue
                if cat not in combined[ph]:
                    combined[ph][cat] = []
                combined[ph][cat].append(('new', st))

    # Write output
    doc = Document()

    for ph, cats in combined.items():
        p_phase = doc.add_paragraph(ph)
        for run in p_phase.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0, 176, 80)   # #00B050

        for cat, items in cats.items():
            p_cat = doc.add_paragraph(cat)
            for run in p_cat.runs:
                run.bold = True
            # Sort items by year desc within category
            items_sorted = sorted(items, key=lambda it: int(it[1].year) if it[1].year.isdigit() else -1, reverse=True)

            for tag, st in items_sorted:
                p = doc.add_paragraph()
                set_hanging_indent_with_tab(p, indent_inch=indent_inch)
                r_year = p.add_run(st.year)
                r_year.bold = False
                p.add_run('\t')
                # For both old and new, write runs; for new we bold until colon
                write_runs(p, st.runs, bold_until_colon=True)

    doc.save(out_docx)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Merge existing sorted .docx with NEW studies (> latest year) from master .docx into one .docx')
    parser.add_argument('--existing-docx', required=True)
    parser.add_argument('--master-docx', required=True)
    parser.add_argument('--out-docx', required=True)
    parser.add_argument('--indent', type=float, default=0.5)

    args = parser.parse_args()

    if not os.path.isfile(args.existing_docx):
        print(f'ERROR: Existing sorted .docx not found: {args.existing_docx}')
        raise SystemExit(2)
    if not os.path.isfile(args.master_docx):
        print(f'ERROR: Master .docx not found: {args.master_docx}')
        raise SystemExit(2)

    try:
        merge_write(args.existing_docx, args.master_docx, args.out_docx, indent_inch=args.indent)
        print('Done.')
        print(f'  Output written to: {os.path.abspath(args.out_docx)}')
    except Exception as e:
        print(f'ERROR: {e}')
        raise SystemExit(1)
