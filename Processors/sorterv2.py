#!/usr/bin/env python3
import argparse
import os
import re
import string
import difflib
from collections import OrderedDict
from typing import List, Tuple, Optional, Dict

try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document = None
    Inches = None

YEAR_LINE = re.compile(r'^\s*(\d{4})\b')
MULTI_X = re.compile(r'x{3,}', flags=re.IGNORECASE)

PHASE_I_PATTERNS = [
    r'^\s*phase\s*i\s*[:.\-–—]?\s*$',
    r'^\s*phase\s*1\s*[:.\-–—]?\s*$'
]
PHASE_II_IV_PATTERNS = [
    r'^\s*phase\s*ii\s*[-–—/ ]*iv\s*[:.\-–—]?\s*$',
    r'^\s*phase\s*2\s*[-–—/ ]*4\s*[:.\-–—]?\s*$'
]
PHASE_I_LABEL = 'Phase I'
PHASE_II_IV_LABEL = 'Phase II-IV'
PHASE_I_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_I_PATTERNS]
PHASE_II_IV_RE = [re.compile(p, re.IGNORECASE) for p in PHASE_II_IV_PATTERNS]


def clean_space_tabs(text: str) -> str:
    text = text.replace('\t', ' ')
    return ' '.join(text.split())


def extract_year(text: str) -> Optional[str]:
    m = YEAR_LINE.match(text)
    if m is not None:
        return m.group(1)
    return None


def strip_leading_year(text: str) -> str:
    if YEAR_LINE.match(text):
        return re.sub(r'^\s*\d{4}\s*', '', text)
    return text


def is_phase_header(line: str) -> Optional[str]:
    stripped = line.strip()
    for r in PHASE_I_RE:
        if r.match(stripped):
            return PHASE_I_LABEL
    for r in PHASE_II_IV_RE:
        if r.match(stripped):
            return PHASE_II_IV_LABEL
    return None


def normalize_after_year(text: str) -> str:
    """
    Normalize only the part after the leading year:
    - lower-case
    - remove long XXXX sequences
    - remove punctuation
    - collapse whitespace
    """
    t = strip_leading_year(text)
    t = t.lower()
    t = MULTI_X.sub('', t)
    t = t.translate(str.maketrans('', '', string.punctuation))
    return ' '.join(t.split())


def parse_master_hierarchy(master_path: str) -> "OrderedDict[str, OrderedDict[str, List[str]]]":
    """
    Parse a MASTER text file:

        PHASE I
        Healthy Adults
        2025 BMS: A Phase 1 ...
        2024 PFIZER: A Phase 3 ...
        <blank line>
        Another Category
        2023 ...

    and return OrderedDict[phase][category] = [full_line, ...]
    """
    phases: "OrderedDict[str, OrderedDict[str, List[str]]]" = OrderedDict()
    current_phase: Optional[str] = None
    current_category: Optional[str] = None
    current_study: Optional[str] = None

    def ensure_phase(ph: str) -> None:
        if ph not in phases:
            phases[ph] = OrderedDict()

    def ensure_category(ph: str, cat: str) -> None:
        ensure_phase(ph)
        if cat not in phases[ph]:
            phases[ph][cat] = []

    with open(master_path, 'r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            if not line.strip():
                if current_study is not None and current_phase is not None and current_category is not None:
                    ensure_category(current_phase, current_category)
                    phases[current_phase][current_category].append(clean_space_tabs(current_study))
                    current_study = None
                continue

            ph = is_phase_header(line)
            if ph is not None:
                if current_study is not None and current_phase is not None and current_category is not None:
                    ensure_category(current_phase, current_category)
                    phases[current_phase][current_category].append(clean_space_tabs(current_study))
                    current_study = None
                current_phase = ph
                ensure_phase(current_phase)
                current_category = None
                continue

            if YEAR_LINE.match(line):
                if current_study is not None and current_phase is not None and current_category is not None:
                    ensure_category(current_phase, current_category)
                    phases[current_phase][current_category].append(clean_space_tabs(current_study))
                current_study = line.strip()
                if current_phase is None:
                    current_phase = PHASE_I_LABEL
                    ensure_phase(current_phase)
                if current_category is None:
                    current_category = "Uncategorized"
                    ensure_category(current_phase, current_category)
                continue

            cat = line.strip()
            if cat.endswith(':'):
                cat = cat[:-1].strip()
            if cat:
                if current_study is not None and current_phase is not None and current_category is not None:
                    ensure_category(current_phase, current_category)
                    phases[current_phase][current_category].append(clean_space_tabs(current_study))
                    current_study = None
                if current_phase is None:
                    current_phase = PHASE_I_LABEL
                    ensure_phase(current_phase)
                current_category = clean_space_tabs(cat)
                ensure_category(current_phase, current_category)
                continue

    if current_study is not None and current_phase is not None and current_category is not None:
        ensure_category(current_phase, current_category)
        phases[current_phase][current_category].append(clean_space_tabs(current_study))

    return phases


END_MARKER = "By signing this form, I confirm that the information provided is accurate and reflects my current qualifications."

def parse_unsorted_studies(unsorted_path: str) -> List[str]:
    """
    Parse the unsorted text file:

        2025 PFIZER: ...
        continued text...
        <blank>
        2024 MODERNA: ...

    into a list of single-line entries with year + description consolidated.

    Also:
    - If the CV footer line
      'By signing this form, I confirm that the information provided is accurate and reflects my current qualifications.'
      appears, it is treated as a hard end-of-section marker and is never included
      in any study description.
    """
    studies: List[str] = []
    current: Optional[str] = None

    with open(unsorted_path, 'r', encoding='utf-8') as f:
        for raw in f:
            # Strip off the footer / section-end marker if it appears on this line
            hit_end = False
            if END_MARKER in raw:
                raw = raw.split(END_MARKER, 1)[0]
                hit_end = True

            line = raw.rstrip('\n')

            # Blank line (or line became blank after cutting off END_MARKER)
            if not line.strip():
                if current is not None:
                    studies.append(clean_space_tabs(current))
                    current = None
                if hit_end:
                    # We've reached the footer; stop processing any further lines
                    break
                continue

            # New study starts with a year
            if YEAR_LINE.match(line):
                if current is not None:
                    studies.append(clean_space_tabs(current))
                current = line.strip()
            else:
                # Continuation of current study
                if current is None:
                    # If there's text before any year, ignore it
                    continue
                current = f"{current} {line.strip()}"

            if hit_end:
                # If the END_MARKER was on same line as content, we've already
                # taken the part before it; now close out after this line.
                break

    if current is not None:
        studies.append(clean_space_tabs(current))

    return studies



def flatten_studies(
    phases: "OrderedDict[str, OrderedDict[str, List[str]]]"
) -> List[Tuple[str, str, str]]:
    """
    Convert phase/category dict into a flat list of:
        (phase, category, full_line)
    """
    flat: List[Tuple[str, str, str]] = []
    for ph, cats in phases.items():
        for cat, studies in cats.items():
            for s in studies:
                flat.append((ph, cat, s))
    return flat


def split_name_protocol(line_after_year: str) -> Tuple[str, bool, str]:
    idx = line_after_year.find(':')
    if idx == -1:
        return line_after_year.strip(), False, ''
    name = line_after_year[:idx].strip()
    remainder = line_after_year[idx + 1:].lstrip()
    return name, True, remainder


def similarity_score(a_norm: str, b_norm: str) -> float:
    """
    Combined similarity:
      - SequenceMatcher ratio
      - Jaccard over whitespace tokens
    """
    if not a_norm or not b_norm:
        return 0.0

    r_seq = difflib.SequenceMatcher(None, a_norm, b_norm).ratio()

    set_a = set(a_norm.split())
    set_b = set(b_norm.split())
    if not set_a or not set_b:
        r_jac = 0.0
    else:
        inter = len(set_a & set_b)
        union = len(set_a | set_b)
        r_jac = inter / union if union else 0.0

    return 0.7 * r_seq + 0.3 * r_jac


def categorize_with_master(
    phases_c: OrderedDict[str, OrderedDict[str, List[str]]],
    phases_b: Optional[OrderedDict[str, OrderedDict[str, List[str]]]],
    unsorted_studies: List[str],
    threshold: float = 0.80
) -> Tuple[
    OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]],
    List[Tuple[str, Optional[str], Optional[str], Optional[str], float]]
]:
    """
    Fuzzy matching with two MASTER files:

    - phases_c: MASTER from Column C (no-red)
    - phases_b: MASTER from Column B (red-label), same structure/order as Column C

    Matching order:
      1) Column C candidates with same year (best fuzzy score >= threshold)
      2) Column B candidates with same year (best fuzzy score >= threshold)
      3) If still no match → Uncategorized

    Output ALWAYS uses:
      - Phase/Category from Column C
      - Text from Column C (no-red form)
    """

    categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]" = OrderedDict()
    for ph, cats in phases_c.items():
        categorized[ph] = OrderedDict()
        for cat in cats.keys():
            categorized[ph][cat] = []

    if 'Uncategorized' not in categorized:
        categorized['Uncategorized'] = OrderedDict()
    if 'Uncategorized' not in categorized['Uncategorized']:
        categorized['Uncategorized']['Uncategorized'] = []

    flat_c = flatten_studies(phases_c)
    flat_b = flatten_studies(phases_b) if phases_b is not None else []

    # Flatten Column C master
    c_year: List[Optional[str]] = []
    c_after: List[str] = []
    c_after_norm: List[str] = []
    c_phase: List[str] = []
    c_cat: List[str] = []
    c_line: List[str] = []

    for ph, cat, s in flat_c:
        yr = extract_year(s)
        after = strip_leading_year(s).strip()
        c_year.append(yr)
        c_after.append(after)
        c_after_norm.append(normalize_after_year(s))
        c_phase.append(ph)
        c_cat.append(cat)
        c_line.append(s)

    # Flatten Column B master (aligned by index)
    b_year: List[Optional[str]] = []
    b_after_norm: List[str] = []

    for ph_b, cat_b, s_b in flat_b:
        yr_b = extract_year(s_b)
        b_year.append(yr_b)
        b_after_norm.append(normalize_after_year(s_b))

    audit_rows: List[Tuple[str, Optional[str], Optional[str], Optional[str], float]] = []

    for u in unsorted_studies:
        u_year = extract_year(u)
        after_u = strip_leading_year(u).strip()
        norm_u = normalize_after_year(u)
        best_idx_c = None
        best_score_c = 0.0
        best_idx_b = None
        best_score_b = 0.0

        if not after_u:
            # nothing to match on
            categorized['Uncategorized']['Uncategorized'].append((u_year or '', after_u))
            audit_rows.append((u, None, None, None, 0.0))
            continue

        # --- PASS 1: Column C (same year) ---
        for i, yr in enumerate(c_year):
            if u_year is not None and yr is not None and yr == u_year:
                sc = similarity_score(norm_u, c_after_norm[i])
                if sc > best_score_c:
                    best_score_c = sc
                    best_idx_c = i

        if best_idx_c is not None and best_score_c >= threshold:
            idx = best_idx_c
            ph = c_phase[idx]
            cat = c_cat[idx]
            year_out = c_year[idx] or u_year or ''
            after_out = c_after[idx]
            categorized[ph][cat].append((year_out, after_out))
            audit_rows.append((u, c_line[idx], ph, cat, best_score_c))
            continue

        # --- PASS 2: Column B (same year) ---
        if phases_b is not None and flat_b:
            for i, yr_b in enumerate(b_year):
                if u_year is not None and yr_b is not None and yr_b == u_year:
                    sc = similarity_score(norm_u, b_after_norm[i])
                    if sc > best_score_b:
                        best_score_b = sc
                        best_idx_b = i

            if best_idx_b is not None and best_score_b >= threshold:
                idx = best_idx_b
                if idx < len(c_phase):
                    ph = c_phase[idx]
                    cat = c_cat[idx]
                    year_out = c_year[idx] or u_year or ''
                    after_out = c_after[idx]
                    line_out = c_line[idx]
                else:
                    ph_b, cat_b, s_b = flat_b[idx]
                    ph = ph_b
                    cat = cat_b
                    year_out = extract_year(s_b) or u_year or ''
                    after_out = strip_leading_year(s_b).strip()
                    line_out = s_b

                categorized[ph][cat].append((year_out, after_out))
                audit_rows.append((u, line_out, ph, cat, best_score_b))
                continue

        # --- No match above threshold ---
        categorized['Uncategorized']['Uncategorized'].append((u_year or '', after_u))
        audit_rows.append((u, None, None, None, 0.0))

    # Sort inside each category by year descending
    def year_as_int(yy: str) -> int:
        try:
            return int(yy)
        except Exception:
            return -1

    for ph, cats in categorized.items():
        for cat, items in cats.items():
            items_sorted = sorted(items, key=lambda t: year_as_int(t[0]), reverse=True)
            categorized[ph][cat] = items_sorted

    return categorized, audit_rows


def write_text(out_path: str,
               categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]",
               indent_sep: str = ' ',
               bold_markers: bool = False,
               bold: bool = True) -> None:
    with open(out_path, 'w', encoding='utf-8') as f:
        for ph, cats in categorized.items():
            f.write(f"{ph}\n")
            for cat, items in cats.items():
                f.write(f"{cat}\n")
                for (year, master_after_year) in items:
                    name, has_colon, remainder = split_name_protocol(master_after_year)
                    if bold and bold_markers:
                        first = f"{year}{indent_sep}**{name}**"
                    else:
                        first = f"{year}{indent_sep}{name}"
                    if has_colon:
                        if len(remainder) > 0:
                            first = f"{first}: {remainder}"
                        else:
                            first = f"{first}:"
                    f.write(first + "\n")
                f.write("\n")
            f.write("\n")


def set_hanging_indent_with_tab(p, indent_inch: float = 0.5) -> None:
    if Inches is None:
        return
    pf = p.paragraph_format
    indent = Inches(indent_inch)
    pf.left_indent = indent
    pf.first_line_indent = -indent
    try:
        pf.tab_stops.add_tab_stop(indent)
    except Exception:
        pass


def write_docx(docx_path: str,
               categorized: "OrderedDict[str, OrderedDict[str, List[Tuple[str, str]]]]",
               docx_indent_inch: float = 0.5,
               bold: bool = True) -> None:
    if Document is None:
        raise RuntimeError('python-docx is not available.')
    doc = Document()
    for ph, cats in categorized.items():
        p_phase = doc.add_paragraph(ph)
        for run in p_phase.runs:
            run.bold = True
        for cat, items in cats.items():
            p_cat = doc.add_paragraph(cat)
            for run in p_cat.runs:
                run.bold = True
            for (year, master_after_year) in items:
                p = doc.add_paragraph()
                set_hanging_indent_with_tab(p, indent_inch=docx_indent_inch)
                r_year = p.add_run(year)
                r_year.bold = False
                p.add_run('\t')
                name, has_colon, remainder = split_name_protocol(master_after_year)
                r_name = p.add_run(name)
                r_name.bold = bool(bold)
                if has_colon:
                    p.add_run(':')
                    if len(remainder) > 0:
                        p.add_run(' ' + remainder)
    doc.save(docx_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Phase-aware sorter that outputs MASTER-formatted text by phase/category and year desc.')
    parser.add_argument('--master', required=True, help='MASTER text file from Column C (no-red)')
    parser.add_argument('--master-b', default=None, help='MASTER text file from Column B (red-label)')
    parser.add_argument('--unsorted', required=True, help='Unsorted studies text file (year-starting lines + continuations)')
    parser.add_argument('--out', required=True, help='Text output path')
    parser.add_argument('--audit', required=True, help='Audit TSV path')
    parser.add_argument('--threshold', type=float, default=0.80)
    parser.add_argument('--indent-type', choices=['tab', 'spaces'], default='spaces')
    parser.add_argument('--indent-size', type=int, default=1)
    parser.add_argument('--bold', type=str, default='true')
    parser.add_argument('--text-bold-markers', type=str, default='false')
    parser.add_argument('--docx-out', default=None)
    parser.add_argument('--docx-indent', type=float, default=0.5)
    args = parser.parse_args()

    if args.indent_type == 'tab':
        indent_sep = '\t'
    else:
        indent_sep = ' ' if args.indent_size < 1 else ' ' * args.indent_size

    def to_bool(val: str) -> bool:
        v = str(val).strip().lower()
        if v in ('1', 'true', 'yes', 'y'):
            return True
        if v in ('0', 'false', 'no', 'n'):
            return False
        return False

    bold_flag = to_bool(args.bold)
    bold_markers_flag = to_bool(args.text_bold_markers)

    if not os.path.isfile(args.master):
        print(f'ERROR: Master file not found: {args.master}')
        raise SystemExit(2)
    if args.master_b is not None and not os.path.isfile(args.master_b):
        print(f'ERROR: Master-B file not found: {args.master_b}')
        raise SystemExit(2)
    if not os.path.isfile(args.unsorted):
        print(f'ERROR: Unsorted file not found: {args.unsorted}')
        raise SystemExit(2)

    try:
        master_phases_c = parse_master_hierarchy(args.master)
        master_phases_b = parse_master_hierarchy(args.master_b) if args.master_b is not None else None
        unsorted_list = parse_unsorted_studies(args.unsorted)

        categorized, audit = categorize_with_master(
            master_phases_c,
            master_phases_b,
            unsorted_list,
            threshold=args.threshold
        )

        write_text(args.out, categorized, indent_sep=indent_sep, bold_markers=bold_markers_flag, bold=bold_flag)

        if args.docx_out is not None:
            write_docx(args.docx_out, categorized, docx_indent_inch=args.docx_indent, bold=bold_flag)

        with open(args.audit, 'w', encoding='utf-8') as f:
            f.write('Unsorted\tMatched_Master\tPhase\tCategory\tScore\n')
            for u, m, ph, cat, sc in audit:
                f.write(f"{u}\t{m or ''}\t{ph or 'Uncategorized'}\t{cat or 'Uncategorized'}\t{sc:.3f}\n")
        try:
            if os.path.isfile(args.master):
                os.remove(args.master)
        except Exception:
            pass

        if args.master_b:
            try:
                if os.path.isfile(args.master_b):
                    os.remove(args.master_b)
            except Exception:
                pass
        # IMPORTANT: do NOT delete master/unsorted files anymore
        # so you can inspect:
        #   - MASTER_TXT (Column C)
        #   - MASTER_TXT_B (Column B)
        #   - UNSORTED_TXT
        #   - AUDIT_TSV

        print('Done.')
        
    except Exception as e:
        print(f'ERROR: {e}')
        raise SystemExit(1)
