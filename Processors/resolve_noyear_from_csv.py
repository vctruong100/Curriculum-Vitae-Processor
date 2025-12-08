
#!/usr/bin/env python3
"""
resolve_noyear_from_csv.py

Goal
----
Some CVs contain study lines under "Research Experience (...)" that DO NOT start with a year.
These lines are already under sponsor categories (e.g., ABBVIE, AMGEN) but lack "YYYY PROTOCOL: ...".
This tool:
  1) Scans the CV's Research Experience section for candidate "no-year" study lines.
  2) Fuzzy matches each line against a mapping CSV (col A=Year, col C=Non-Red study text).
  3) If above threshold, it produces lines like "YYYY <non-red study text from CSV>" so they will
     match your MASTER .txt in the sorter.
  4) Merges those resolved lines into your existing ".ADD_CV_STUDIES_FROM_DOCX.txt" (unsorted list) de-duplicated.

CSV format (same expectation as your remove-red tool):
  Column A: Year           (e.g., 2021)
  Column B: Red study text (ignored here)
  Column C: Non-red study text (used for fuzzy comparison + output description)

Usage
-----
python resolve_noyear_from_csv.py \
  --cv "Full_CV.docx" \
  --csv "mapping.csv" \
  --in-unsorted "Output/.ADD_CV_STUDIES_FROM_DOCX.txt" \
  --out-unsorted "Output/.ADD_CV_STUDIES_FROM_DOCX.txt" \
  --section-start "Research Experience" \
  --section-end   "By signing this form, I confirm that the information provided is accurate and reflects my current qualifications." \
  --threshold 0.88 \
  --audit "Output/noyear_resolve_audit.tsv"

Notes
-----
- Requires: python-docx
- Encoding: tries utf-8-sig then latin-1 for the CSV (robust to BOM and Windows-1252 sources).
- We intentionally keep filenames unchanged. The default behavior is to UPDATE the same unsorted .txt in place.
"""

import argparse
import csv
import os
import re
import string
from typing import List, Tuple, Dict

try:
    from docx import Document
except Exception:
    Document = None

YEAR_RE = re.compile(r'^\s*(\d{4})\b')

def _norm_ws(s: str) -> str:
    return ' '.join((s or '').replace('\t',' ').split())

def _normalize_for_match(s: str) -> str:
    # Lower, remove punctuation (but keep spaces), collapse spaces
    s = (s or '').lower()
    s = s.translate(str.maketrans({c:' ' for c in string.punctuation}))
    return _norm_ws(s)

def _similarity(a: str, b: str) -> float:
    try:
        import difflib
        return difflib.SequenceMatcher(None, _normalize_for_match(a), _normalize_for_match(b)).ratio()
    except Exception:
        A = set(_normalize_for_match(a).split())
        B = set(_normalize_for_match(b).split())
        if not A and not B: return 1.0
        if not A or not B:  return 0.0
        return len(A & B) / len(A | B)

def _load_csv_mapping(csv_path: str) -> List[Tuple[str, str]]:
    """
    Returns list of (year, nonred_text_after_year). We purposely ignore the red-text column.
    Robust to encodings (utf-8-sig first, then latin-1).
    """
    encs = ['utf-8-sig', 'latin-1']
    last_err = None
    rows: List[Tuple[str, str]] = []
    for enc in encs:
        try:
            with open(csv_path, 'r', encoding=enc, newline='') as f:
                rdr = csv.reader(f)
                for row in rdr:
                    if not row or len(row) < 3:
                        continue
                    year = str(row[0]).strip()
                    if not (year.isdigit() and len(year) == 4):
                        # header or invalid
                        continue
                    nonred = str(row[2]).strip()
                    if not nonred:
                        continue
                    rows.append((year, _norm_ws(nonred)))
            return rows
        except UnicodeDecodeError as e:
            last_err = e
            continue
    raise RuntimeError(f"Could not read CSV with utf-8 or latin-1: {last_err}")

def _find_bounds(doc: "Document", start_text: str, end_text: str) -> Tuple[int, int]:
    def norm(s): return ' '.join((s or '').strip().lower().split())
    s_norm = norm(start_text); e_norm = norm(end_text)
    paras = list(doc.paragraphs)

    s_idx = None
    for i, p in enumerate(paras):
        if s_norm in norm(p.text):
            s_idx = i
            break
    if s_idx is None:
        raise RuntimeError(f'Could not find section start containing: "{start_text}".')
    e_idx = None
    for j in range(s_idx + 1, len(paras)):
        if e_norm in norm(paras[j].text):
            e_idx = j
            break
    if e_idx is None:
        e_idx = len(paras)
    return s_idx, e_idx

def _collect_noyear_candidates(doc: "Document", s_idx: int, e_idx: int) -> List[str]:
    """
    Heuristics:
      - Must be non-empty
      - Must NOT start with a year
      - Likely a study line if it has a comma (role, descriptors) OR length >= 40 characters
      - Category headers (ABBVIE, AMGEN, etc.) will often be short (filtered by length/comma)
    """
    paras = list(doc.paragraphs)
    cands: List[str] = []
    for i in range(s_idx + 1, e_idx):
        t = (paras[i].text or '').strip()
        if not t:
            continue
        if YEAR_RE.match(t):
            continue
        # Likely category header? Skip ultra short and no comma.
        if (len(t) < 40) and (',' not in t):
            continue
        cands.append(_norm_ws(t))
    return cands

def _dedupe_keep_order(lines: List[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for s in lines:
        k = _normalize_for_match(s)
        if k in seen:
            continue
        seen.add(k)
        out.append(s)
    return out

def resolve_and_merge(cv_path: str, csv_path: str, in_unsorted: str, out_unsorted: str,
                      section_start: str, section_end: str, threshold: float, audit_path: str = None) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required.")
    if not os.path.isfile(cv_path):
        raise FileNotFoundError(f"CV not found: {cv_path}")
    if not os.path.isfile(csv_path):
        raise FileNotFoundError(f"CSV not found: {csv_path}")
    if not os.path.isfile(in_unsorted):
        raise FileNotFoundError(f"Unsorted .txt not found: {in_unsorted}")

    doc = Document(cv_path)
    s_idx, e_idx = _find_bounds(doc, section_start, section_end)
    cands = _collect_noyear_candidates(doc, s_idx, e_idx)
    mapping = _load_csv_mapping(csv_path)

    resolved: List[str] = []
    audit_rows: List[Tuple[str, str, float]] = []  # (cand, chosen_nonred, score)

    # Build a searchable index of mapping lines
    nonred_only = [nr for (_, nr) in mapping]

    import difflib
    for cand in cands:
        best = None
        best_score = -1.0
        for (year, nonred) in mapping:
            sc = _similarity(cand, nonred)
            if sc > best_score:
                best_score = sc
                best = (year, nonred)
        if best and best_score >= threshold:
            year, nonred = best
            resolved.append(f"{year} {nonred}")
            audit_rows.append((cand, nonred, best_score))

    # Read current unsorted lines
    with open(in_unsorted, 'r', encoding='utf-8') as f:
        base_lines = [ln.strip() for ln in f.read().splitlines() if ln.strip()]

    merged = _dedupe_keep_order(base_lines + resolved)

    # Write back to the same out_unsorted (can be the same as in_unsorted)
    with open(out_unsorted, 'w', encoding='utf-8') as f:
        for ln in merged:
            f.write(ln + "\n\n")

    if audit_path:
        with open(audit_path, 'w', encoding='utf-8', newline='') as f:
            f.write("candidate\tmatched_nonred\tscore\n")
            for cand, nonred, sc in audit_rows:
                f.write(f"{cand}\t{nonred}\t{sc:.4f}\n")

def main():
    ap = argparse.ArgumentParser(description="Resolve no-year study lines in a CV using CSV fuzzy mapping and merge into unsorted .txt")
    ap.add_argument('--cv', required=True, help='Original CV .docx')
    ap.add_argument('--csv', required=True, help='Mapping CSV (A=year, C=non-red text)')
    ap.add_argument('--in-unsorted', required=True, help='Existing unsorted .txt (from extract step)')
    ap.add_argument('--out-unsorted', required=True, help='Destination unsorted .txt (often same as --in-unsorted)')
    ap.add_argument('--section-start', default="Research Experience")
    ap.add_argument('--section-end', default="By signing this form, I confirm that the information provided is accurate and reflects my current qualifications.")
    ap.add_argument('--threshold', type=float, default=0.88, help='Fuzzy match threshold (0.80–0.98 typical)')
    ap.add_argument('--audit', default=None, help='Optional TSV audit output')
    args = ap.parse_args()

    resolve_and_merge(args.cv, args.csv, args.in_unsorted, args.out_unsorted,
                      args.section_start, args.section_end, args.threshold, args.audit)

if __name__ == "__main__":
    main()
