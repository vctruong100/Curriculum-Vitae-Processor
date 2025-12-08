
#!/usr/bin/env python3
"""
inject_sorted_into_cv.py  (enhanced v2: preserve Phase headers)

- Clones the original CV into Output.
- Replaces ONLY the "Research Experience" section body (until the disclaimer).
- Injects formatted studies from ".UPDATED CV.docx" (preserves bold/italic/underline,
  font name/size, and **font color** — including green Phase headers and red labels).
- After injection, removes any category that has no studies under it,
  **except** Phase headers (e.g., "Phase I", "Phase II", etc.) which are always kept.
"""

import argparse
import shutil
import os
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

YEAR_RE = re.compile(r'^\s*(\d{4})\b')
PHASE_RE = re.compile(r'^\s*phase(\s|$)', re.IGNORECASE)

DEFAULT_SECTION_START = "Research Experience"
DEFAULT_SECTION_END = "By signing this form, I confirm that the information provided is accurate and reflects my current qualifications."

def _add_top_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    # remove existing pBdr if any
    for child in list(pPr):
        if child.tag == qn('w:pBdr'):
            pPr.remove(child)
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')   
    top.set(qn('w:sz'),  '14')        # thickness
    top.set(qn('w:space'),'1')       # spacing
    top.set(qn('w:color'),'auto')
    pBdr.append(top)
    pPr.append(pBdr)


def _remove_top_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    removed = False
    for child in list(pPr):
        if child.tag == qn('w:pBdr'):
            pPr.remove(child)
            removed = True
    return removed


def _norm(s: str) -> str:
    return ' '.join((s or '').strip().lower().split())

def _all_body_paragraphs(doc):
    for p in doc.paragraphs:
        yield p

def _find_bounds_paras(doc, start_text: str, end_text: str):
    paras = list(_all_body_paragraphs(doc))
    s_norm = _norm(start_text)
    e_norm = _norm(end_text)

    s_idx = None
    for i, p in enumerate(paras):
        if s_norm in _norm(p.text):
            s_idx = i
            break
    if s_idx is None:
        raise RuntimeError(f'Could not find section start containing: "{start_text}".')

    e_idx = None
    for j in range(s_idx + 1, len(paras)):
        if e_norm in _norm(paras[j].text):
            e_idx = j
            break
    if e_idx is None:
        e_idx = len(paras)

    return paras, s_idx, e_idx

def _delete_paragraph(p):
    p._element.getparent().remove(p._element)

def _insert_paragraph_after(reference_paragraph):
    new_p = OxmlElement('w:p')
    reference_paragraph._p.addnext(new_p)
    return Paragraph(new_p, reference_paragraph._parent)

def _copy_run_format(src_run, dst_run):
    rfont = dst_run.font
    sfont = src_run.font
    rfont.bold = sfont.bold
    rfont.italic = sfont.italic
    rfont.underline = sfont.underline
    # size / name
    if sfont.size is not None:
        rfont.size = sfont.size
    if sfont.name is not None:
        rfont.name = sfont.name
    # color (preserve red/green etc.)
    try:
        scol = sfont.color
        if getattr(scol, "rgb", None) is not None:
            rfont.color.rgb = scol.rgb
        elif getattr(scol, "theme_color", None) is not None:
            rfont.color.theme_color = scol.theme_color
        elif getattr(scol, "type", None) is not None:
            rfont.color._element = scol._element
    except Exception:
        pass

def _copy_paragraph_format(src_p, dst_p):
    dst_p.alignment = src_p.alignment
    pf = dst_p.paragraph_format
    spf = src_p.paragraph_format
    pf.left_indent = spf.left_indent
    pf.first_line_indent = spf.first_line_indent
    pf.space_before = spf.space_before
    pf.space_after = spf.space_after
    pf.keep_together = spf.keep_together
    pf.keep_with_next = spf.keep_with_next
    pf.widow_control = spf.widow_control
    try:
        if src_p.style is not None and src_p.style.name:
            dst_p.style = src_p.style
    except Exception:
        pass

def _append_paragraph_like(dst_after_p, src_p):
    new_p = _insert_paragraph_after(dst_after_p)
    _copy_paragraph_format(src_p, new_p)
    if src_p.runs:
        for r in src_p.runs:
            nr = new_p.add_run(r.text)
            _copy_run_format(r, nr)
    else:
        new_p.add_run("")
    return new_p

def _is_study_paragraph(p):
    return YEAR_RE.match(p.text or '') is not None

def _is_category_paragraph(p):
    txt = (p.text or '').strip()
    if not txt:
        return False
    if _is_study_paragraph(p):
        return False
    try:
        if p.style and 'heading' in p.style.name.lower():
            return True
    except Exception:
        pass
    # No digits, reasonably short
    if any(ch.isdigit() for ch in txt):
        return False
    if len(txt) <= 80:
        ends_colon = txt.endswith(':')
        mostly_upper = sum(c.isupper() for c in txt if c.isalpha()) >= 0.6 * sum(1 for c in txt if c.isalpha())
        all_bold = (len(p.runs) > 0) and all((r.font.bold is True) for r in p.runs if r.text.strip())
        if ends_colon or mostly_upper or all_bold:
            return True
    return False

def _prune_empty_categories(out_doc, section_start, section_end):
    """Remove category paragraphs in the Research Experience block that have no studies
       until the next category or the disclaimer, but NEVER remove Phase headers.
    """
    paras, s_idx, e_idx = _find_bounds_paras(out_doc, section_start, section_end)
    start = s_idx + 1
    end = e_idx - 1
    if start > end:
        return

    block = paras[start:e_idx]

    def g(i): return start + i

    cat_positions = []
    for i, p in enumerate(block):
        if (start + i) >= e_idx:
            break
        if _is_category_paragraph(p):
            cat_positions.append(i)

    to_delete_global_idxs = []
    for idx, cat_i in enumerate(cat_positions):
        p_cat = block[cat_i]
        txt = (p_cat.text or '').strip()
        # Keep Phase headers unconditionally
        if PHASE_RE.match(txt):
            continue
        scan_from = cat_i + 1
        scan_to = cat_positions[idx + 1] if idx + 1 < len(cat_positions) else (e_idx - start)
        has_study = False
        for j in range(scan_from, scan_to):
            pj = block[j]
            if _is_study_paragraph(pj):
                has_study = True
                break
        if not has_study:
            to_delete_global_idxs.append(g(cat_i))

    if to_delete_global_idxs:
        for di in sorted(to_delete_global_idxs, reverse=True):
            _delete_paragraph(paras[di])

def inject_sorted(original_cv: str, studies_docx: str, out_cv: str,
                  section_start: str = DEFAULT_SECTION_START,
                  section_end: str = DEFAULT_SECTION_END):
    from docx import Document

    if not os.path.isfile(original_cv):
        raise FileNotFoundError(f"Original CV not found: {original_cv}")
    if not os.path.isfile(studies_docx):
        raise FileNotFoundError(f"Studies .docx not found: {studies_docx}")

    try:
        shutil.copyfile(original_cv, out_cv)
    except Exception:
        raise PermissionError(f"Permission denied. Please close ALL CV documents.")

    out_doc = Document(out_cv)

    paras, s_idx, e_idx = _find_bounds_paras(out_doc, section_start, section_end)

    for i in range(e_idx - 1, s_idx, -1):
        _delete_paragraph(paras[i])

    out_doc.save(out_cv)
    out_doc = Document(out_cv)
    paras, s_idx2, e_idx2 = _find_bounds_paras(out_doc, section_start, section_end)
    start_p = paras[s_idx2]

    src_doc = Document(studies_docx)
    src_paras = list(src_doc.paragraphs)

    insert_after = start_p
    for sp in src_paras:
        if not sp.text and not sp.runs:
            continue
        insert_after = _append_paragraph_like(insert_after, sp)

    # Ensure there's a blank paragraph after the inserted block (nice spacing)
    blank_p = OxmlElement("w:p")
    insert_after._p.addnext(blank_p)
    insert_after = Paragraph(blank_p, insert_after._parent)

    # Find **all** occurrences of the disclaimer and ensure only the LAST one has the top border.
    end_norm = _norm(section_end)
    matches = [p for p in out_doc.paragraphs if end_norm in _norm(p.text)]
    if matches:
        # Remove any border from all occurrences first
        for p_d in matches:
            try:
                _remove_top_border(p_d)
            except Exception:
                pass
        # Add border to the LAST occurrence (end of full CV)
        _add_top_border(matches[-1])
    else:
        # Don't crash – just continue without the border
        print("WARN: Disclaimer text not found in document; skipping border.")

    _prune_empty_categories(out_doc, section_start, section_end)

    out_doc.save(out_cv)

def main():
    ap = argparse.ArgumentParser(description="Clone a CV and inject sorted studies into Research Experience section.")
    ap.add_argument('--original-cv', required=True, help='Path to the original full CV .docx')
    ap.add_argument('--studies-docx', required=True, help='Path to the .docx that contains the final sorted studies (e.g., .UPDATED CV.docx)')
    ap.add_argument('--out', required=True, help='Output path for the updated CV copy')
    ap.add_argument('--section-start', default=DEFAULT_SECTION_START, help='Section start marker')
    ap.add_argument('--section-end', default=DEFAULT_SECTION_END, help='Section end marker')
    args = ap.parse_args()

    inject_sorted(args.original_cv, args.studies_docx, args.out, args.section_start, args.section_end)
    print("Done.")
    print(f"  Updated CV saved to: {os.path.abspath(args.out)}")

if __name__ == '__main__':
    main()
